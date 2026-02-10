"""
Word document (DOCX) parser using python-docx.
"""

from .document_parser import DocumentParser
import io


class DocxParser(DocumentParser):
    """Parser for Word documents (.docx)."""
    
    def extract_text(self, file_bytes: bytes, filename: str) -> str:
        """
        Extract text from Word document.
        
        Args:
            file_bytes: DOCX file bytes
            filename: Original filename
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If document is invalid or corrupted
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx not installed. Run: pip install python-docx"
            )
        
        try:
            doc = Document(io.BytesIO(file_bytes))
            text_content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_content.append(row_text)
            
            if not text_content:
                raise ValueError("Could not extract any text from Word document")
            
            return '\n\n'.join(text_content)
            
        except Exception as e:
            if isinstance(e, ValueError):
                raise
            raise ValueError(f"Failed to parse Word document: {str(e)}")
